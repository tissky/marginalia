"""DOCX pipeline (.docx via python-docx).

Extracts paragraphs + headings + table cells in document order, then
hands off to the shared text indexer. The original docx hierarchy
(Heading 1/2/3 styles) is preserved as a heading prefix `# / ## / ###`
so the indexer can produce heading-anchored sections.

read_segment supports paragraph_start / paragraph_end ranges (1-indexed,
inclusive — only counting non-empty rendered blocks), regex pattern
search, and the generic offset/max_chars chunking over the full body.

Images, embedded objects, and footnotes are skipped — for image-heavy
decks the user is better served by exporting to PDF and using the
pdf-with-figures pipeline.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any

from marginalia.pipelines._text_indexer import index_extracted_text
from marginalia.pipelines.base import (
    Pipeline,
    PipelineContext,
    PipelineResult,
    SegmentResult,
)
from marginalia.pipelines.registry import register_pipeline
from marginalia.storage.base import StorageBackend

log = logging.getLogger(__name__)

MAX_DOCX_BYTES = 30 * 1024 * 1024  # 30 MB hard cap
MAX_OUTPUT_CHARS = 80_000  # plenty for the LLM prompt
DEFAULT_MAX_CHARS = 8000


@register_pipeline(
    mimes=(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    exts=(".docx",),
)
class DocxPipeline(Pipeline):
    name = "docx"

    async def run(
        self,
        *,
        ctx: PipelineContext,
        storage: StorageBackend,
    ) -> PipelineResult:
        paragraphs = await self._extract_paragraphs(storage, ctx.storage_key)
        body = "\n".join(paragraphs)
        if len(body) > MAX_OUTPUT_CHARS:
            body = body[:MAX_OUTPUT_CHARS] + "\n[…document truncated for indexing…]"
        return await index_extracted_text(body, ctx, kind="text")

    async def read_segment(
        self,
        *,
        file_row: Any,
        args: dict[str, Any],
        storage: StorageBackend,
    ) -> SegmentResult:
        paragraphs = await self._extract_paragraphs(storage, file_row.storage_key)
        return self._slice(paragraphs, args)

    async def read_segment_from_bytes(
        self,
        body: bytes,
        args: dict[str, Any],
        *,
        filename: str | None = None,
    ) -> SegmentResult:
        """Bytes-first variant — used by ArchivePipeline for member peeks."""
        if len(body) > MAX_DOCX_BYTES:
            return SegmentResult(
                error=f"docx exceeds {MAX_DOCX_BYTES // (1024*1024)}MB cap",
            )
        try:
            paragraphs = self._parse_paragraphs_from_bytes(body)
        except Exception as exc:  # noqa: BLE001 — python-docx surfaces many
            return SegmentResult(error=f"docx parse failed: {exc}")
        return self._slice(paragraphs, args)

    def _slice(
        self, paragraphs: list[str], args: dict[str, Any],
    ) -> SegmentResult:
        """Resolve args against this docx body.

        Field priority:
          1. pattern                       → regex search
          2. paragraph_start/_end          → return paragraph range
          3. (default)                     → offset..offset+max_chars chunk
        """
        body = "\n".join(paragraphs)
        total_paragraphs = len(paragraphs)

        offset = max(0, int(args.get("offset") or 0))
        max_chars = int(args.get("max_chars") or DEFAULT_MAX_CHARS)
        if max_chars <= 0:
            max_chars = DEFAULT_MAX_CHARS

        pattern = (args.get("pattern") or "").strip()
        if pattern:
            scope_paragraphs = paragraphs
            paragraph_offset = 0
            ps_raw = args.get("paragraph_start")
            pe_raw = args.get("paragraph_end")
            if ps_raw or pe_raw:
                try:
                    ps = max(1, int(ps_raw)) if ps_raw else 1
                    pe = int(pe_raw) if pe_raw else len(paragraphs)
                except (TypeError, ValueError):
                    return SegmentResult(error="paragraph_start/end must be integers")
                if pe < ps:
                    return SegmentResult(error="paragraph_end must be >= paragraph_start")
                ps = max(1, min(ps, max(1, len(paragraphs))))
                pe = max(ps, min(pe, len(paragraphs)))
                scope_paragraphs = paragraphs[ps - 1: pe]
                paragraph_offset = ps - 1
            return _docx_pattern_search(
                paragraphs=scope_paragraphs, pattern=pattern,
                context_lines=int(args.get("context_lines") or 2),
                max_matches=int(args.get("max_matches") or 20),
                match_offset=max(0, int(args.get("match_offset") or 0)),
                paragraph_offset=paragraph_offset,
                total_paragraphs_full=len(paragraphs),
            )

        para_start = args.get("paragraph_start")
        para_end = args.get("paragraph_end")
        if para_start:
            try:
                ps = int(para_start)
            except (TypeError, ValueError):
                return SegmentResult(error="paragraph_start must be an integer")
            try:
                pe = int(para_end) if para_end else ps
            except (TypeError, ValueError):
                return SegmentResult(error="paragraph_end must be an integer")
            if total_paragraphs == 0:
                return SegmentResult(error="docx has no paragraphs")
            ps = max(1, min(ps, total_paragraphs))
            pe = max(ps, min(pe, total_paragraphs))
            slab = "\n".join(paragraphs[ps - 1: pe])
            return _clamp(
                slab, offset, max_chars,
                extras={
                    "paragraph_start": ps,
                    "paragraph_end": pe,
                    "total_paragraphs": total_paragraphs,
                },
            )

        # Compute paragraph range from char offset so footnotes can
        # deep-link even when the LLM reads by offset rather than
        # paragraph_start/paragraph_end.
        para_start = body[:offset].count("\n") + 1
        chunk_for_range = body[offset: offset + max_chars]
        para_end = para_start + chunk_for_range.count("\n")
        return _clamp(
            body, offset, max_chars,
            extras={
                "total_paragraphs": total_paragraphs,
                "paragraph_start": para_start,
                "paragraph_end": para_end,
            },
        )

    @classmethod
    async def _extract_paragraphs(
        cls,
        storage: StorageBackend, key: str,
    ) -> list[str]:
        try:
            from docx import Document  # type: ignore  # noqa: F401 — keeps import-error early
        except ImportError as exc:
            raise RuntimeError(
                "docx pipeline needs python-docx; "
                "`pip install python-docx`"
            ) from exc

        buf = bytearray()
        async for chunk in storage.get(key):
            buf.extend(chunk)
            if len(buf) > MAX_DOCX_BYTES:
                raise ValueError(
                    f"docx exceeds {MAX_DOCX_BYTES // (1024*1024)}MB cap"
                )
        return cls._parse_paragraphs_from_bytes(bytes(buf))

    @staticmethod
    def _parse_paragraphs_from_bytes(body: bytes) -> list[str]:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "docx pipeline needs python-docx; "
                "`pip install python-docx`"
            ) from exc
        doc = Document(io.BytesIO(body))
        out: list[str] = []
        for block in _iter_block_items(doc):
            line = _render_block(block)
            if line:
                out.append(line)
        return out


def _iter_block_items(doc: Any):
    """Yield paragraphs and tables in document order.

    python-docx exposes doc.paragraphs and doc.tables as separate lists, so
    walk the underlying body XML to recover order.
    """
    from docx.oxml.ns import qn  # type: ignore
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _render_block(block: Any) -> str:
    cls_name = type(block).__name__
    if cls_name == "Paragraph":
        text = (block.text or "").strip()
        if not text:
            return ""
        style = (getattr(block.style, "name", "") or "").strip()
        if style.startswith("Heading 1"):
            return f"# {text}"
        if style.startswith("Heading 2"):
            return f"## {text}"
        if style.startswith("Heading 3"):
            return f"### {text}"
        if style.startswith("Heading"):
            return f"#### {text}"
        return text
    if cls_name == "Table":
        rows: list[str] = []
        for row in block.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    return ""


# ---- read_segment helpers --------------------------------------------------

def _clamp(
    text: str, offset: int, max_chars: int,
    *, extras: dict[str, Any] | None = None,
) -> SegmentResult:
    extras = dict(extras or {})
    total = len(text)
    chunk = text[offset: offset + max_chars]
    truncated = (offset + len(chunk)) < total
    extras.update({
        "offset": offset,
        "char_count": len(chunk),
        "total_chars": total,
        "truncated": truncated,
    })
    if truncated:
        extras["next_offset"] = offset + len(chunk)
    if not chunk:
        return SegmentResult(text="", error="empty result", extras=extras)
    return SegmentResult(text=chunk, extras=extras)


def _docx_pattern_search(
    *, paragraphs: list[str], pattern: str,
    context_lines: int, max_matches: int,
    match_offset: int = 0, paragraph_offset: int = 0,
    total_paragraphs_full: int | None = None,
) -> SegmentResult:
    try:
        rx = re.compile(pattern, re.IGNORECASE | re.MULTILINE)
    except re.error as exc:
        return SegmentResult(error=f"invalid regex: {exc}")

    full_total = (
        total_paragraphs_full if total_paragraphs_full is not None
        else len(paragraphs)
    )

    all_hits: list[dict[str, Any]] = []
    for i, para in enumerate(paragraphs, start=1):
        if not para:
            continue
        for m in rx.finditer(para):
            s = max(0, i - 1 - context_lines)
            e = min(len(paragraphs), i + context_lines)
            all_hits.append({
                "paragraph": i + paragraph_offset,
                "match": m.group(0)[:200],
                "context": "\n".join(paragraphs[s:e]),
            })

    total = len(all_hits)
    hits = all_hits[match_offset: match_offset + max_matches]
    has_more = (match_offset + len(hits)) < total

    extras: dict[str, Any] = {
        "pattern": pattern,
        "match_count": len(hits),
        "total_matches": total,
        "match_offset": match_offset,
        "has_more": has_more,
        "hits": hits,
        "total_paragraphs": full_total,
    }
    if has_more:
        extras["next_match_offset"] = match_offset + len(hits)

    if not hits:
        if match_offset and total:
            err = f"match_offset {match_offset} exceeds total_matches {total}"
        else:
            err = "no matches"
        return SegmentResult(text="", error=err, extras=extras)

    rendered = "\n\n".join(
        f"[¶{h['paragraph']}] {h['match']}\n  ┊ {h['context']}"
        for h in hits
    )
    return SegmentResult(text=rendered, extras=extras)
